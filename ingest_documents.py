"""
Phase 1 – Document Ingestion and Indexing
=========================================
Responsibilities
  1. Bulk processing of PDF and Word documents from SECP_docs/
  2. Source-domain validation (secp.gov.pk)
  3. Append-only audit logging
  4. Incremental updates via SHA-256 checksums (skip unchanged files)
  5. Scanned-page detection (flagged for manual OCR review)

Output
  data/processed_docs/<doc_id>.json   – structured content + metadata
  data/flagged/<filename>             – symlink / copy for flagged docs
  data/audit_log.jsonl                – one JSON line per ingestion event
  data/checksums.json                 – filename → sha256 map
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import traceback
import urllib.parse
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument

import config

try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

VISION_OCR_AVAILABLE = bool(config.OPENAI_API_KEY)

# Apply Tesseract path from config as soon as the module loads
if OCR_AVAILABLE and config.TESSERACT_CMD:
    import os as _os
    if _os.path.exists(config.TESSERACT_CMD):
        pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_CMD


# ══════════════════════════════════════════════════════════════════════════════
# Data models
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class PageContent:
    page_num: int
    text: str
    char_count: int
    is_scanned: bool   # True when char_count < threshold


@dataclass
class OrderMetadata:
    """Structured metadata extracted from the SECP filename convention."""
    sections: list[str]          # e.g. ["510"] or ["134", "183", "510"]
    company_name: str
    order_date_raw: str          # as it appears in the filename
    order_date_iso: Optional[str] = None   # "YYYY-MM-DD" when parseable
    act_reference: Optional[str] = None   # e.g. "Companies Act 2017"


@dataclass
class ProcessedDocument:
    doc_id: str                  # SHA-256 of file bytes
    filename: str
    file_path: str
    file_type: str               # "pdf" | "docx" | "doc"
    source_url: Optional[str]    # provided by caller; None = not verified
    source_verified: bool        # True when URL is from secp.gov.pk
    ingestion_timestamp: str
    status: str                  # success | partial | scanned | failed
    order_metadata: OrderMetadata
    pages: list[PageContent]
    full_text: str
    page_count: int
    scanned_page_count: int
    error: Optional[str] = None
    processing_notes: list[str] = field(default_factory=list)


# ══════════════════════════════════════════════════════════════════════════════
# Filename parser
# ══════════════════════════════════════════════════════════════════════════════

class FilenameParser:
    """
    Parse SECP adjudication-order filenames.

    Supported patterns
    ──────────────────
    Standard  : Order-510-Company-Name-31.7.25.pdf
    Multi-sec : Order-134-183-510-Company-Name-24.10.25.pdf
    Hyphen-dt : Order-249-Company-Name-22-5-2025.pdf
    Dated     : Order-dated-07-Nov-25-Company-Name-Sec510-Companies-Act-2017.pdf
    """

    # Date patterns tried in order of specificity
    _DOT_DATE    = re.compile(r'(\d{1,2})\.(\d{1,2})\.(\d{2,4})$')
    _HYPHEN_DATE = re.compile(r'(\d{1,2})-(\d{1,2})-(\d{2,4})$')
    _NAMED_MONTH = re.compile(r'(\d{2})-([A-Za-z]{3})-(\d{2,4})')

    _MONTH_MAP = {
        'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
    }

    @staticmethod
    def _normalise_year(y: str) -> int:
        y = int(y)
        return 2000 + y if y < 100 else y

    @classmethod
    def _to_iso(cls, day: int, month: int, year_raw: str) -> str:
        year = cls._normalise_year(year_raw)
        try:
            return f"{year:04d}-{month:02d}-{day:02d}"
        except Exception:
            return ""

    @classmethod
    def _parse_dated_pattern(cls, stem: str) -> OrderMetadata:
        """Handle  Order-dated-DD-Mon-YY-Company-SecNNN-Companies-Act-YYYY"""
        # Extract named-month date
        m = cls._NAMED_MONTH.search(stem)
        date_raw = m.group(0) if m else ""
        iso_date = ""
        if m:
            day   = int(m.group(1))
            month = cls._MONTH_MAP.get(m.group(2).lower(), 0)
            iso_date = cls._to_iso(day, month, m.group(3))

        # Extract section numbers (SecNNN pattern)
        sections = re.findall(r'[Ss]ec(\d+)', stem)

        # Act reference
        act_ref = "Companies Act 2017" if "Companies-Act-2017" in stem else None

        # Company name: between the date and the Sec marker
        # Remove known noise
        company_raw = stem
        for pattern in [r'^Order-dated-', r'\d{2}-[A-Za-z]{3}-\d{2,4}-?',
                        r'-?[Ss]ec\d+-Companies-Act-\d+']:
            company_raw = re.sub(pattern, '', company_raw)
        company = company_raw.strip('-').replace('-', ' ').strip()

        return OrderMetadata(
            sections=sections,
            company_name=company,
            order_date_raw=date_raw,
            order_date_iso=iso_date or None,
            act_reference=act_ref,
        )

    @classmethod
    def parse(cls, filename: str) -> OrderMetadata:
        stem = Path(filename).stem   # remove extension

        # ── Special "dated" pattern ───────────────────────────────────────────
        if stem.lower().startswith('order-dated-'):
            return cls._parse_dated_pattern(stem)

        # ── Standard pattern ──────────────────────────────────────────────────
        # Remove "Order-" prefix
        body = re.sub(r'^[Oo]rder-', '', stem)

        # 1. Detect and strip date suffix
        date_raw = ""
        iso_date = ""

        m_dot = cls._DOT_DATE.search(body)
        m_hyp = cls._HYPHEN_DATE.search(body)

        if m_dot:
            date_raw = m_dot.group(0)
            iso_date = cls._to_iso(int(m_dot.group(1)), int(m_dot.group(2)), m_dot.group(3))
            body = body[:m_dot.start()].rstrip('-')
        elif m_hyp:
            date_raw = m_hyp.group(0)
            iso_date = cls._to_iso(int(m_hyp.group(1)), int(m_hyp.group(2)), m_hyp.group(3))
            body = body[:m_hyp.start()].rstrip('-')

        # 2. Extract leading section numbers (pure digit tokens)
        parts = body.split('-')
        sections = []
        company_parts = []
        for i, part in enumerate(parts):
            if part.isdigit():
                sections.append(part)
            else:
                company_parts = parts[i:]
                break

        company = ' '.join(company_parts).replace('  ', ' ').strip()
        if not company:
            company = body  # fallback: use full body

        return OrderMetadata(
            sections=sections,
            company_name=company,
            order_date_raw=date_raw,
            order_date_iso=iso_date or None,
        )


# ══════════════════════════════════════════════════════════════════════════════
# Document validators
# ══════════════════════════════════════════════════════════════════════════════

class DocumentValidator:
    """Validates format and, when a URL is supplied, source domain."""

    @staticmethod
    def validate_extension(path: Path) -> tuple[bool, str]:
        ext = path.suffix.lower()
        if ext not in config.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {ext}"
        return True, ""

    @staticmethod
    def validate_source_url(url: Optional[str]) -> tuple[bool, str]:
        """
        Returns (verified, note).
        verified=True  → URL confirmed from secp.gov.pk
        verified=False → URL absent or from a different domain
        """
        if not url:
            return False, "No source URL provided; origin unverified"
        parsed = urllib.parse.urlparse(url)
        host = parsed.netloc.lower().lstrip("www.")
        if config.VALID_SOURCE_DOMAIN not in host:
            return False, f"URL domain '{host}' is not {config.VALID_SOURCE_DOMAIN}"
        return True, ""

    @staticmethod
    def validate_file_exists(path: Path) -> tuple[bool, str]:
        if not path.exists():
            return False, f"File not found: {path}"
        if path.stat().st_size == 0:
            return False, "File is empty (0 bytes)"
        return True, ""


# ══════════════════════════════════════════════════════════════════════════════
# Vision OCR  (GPT-4o)
# ══════════════════════════════════════════════════════════════════════════════

class VisionOCR:
    """
    Uses GPT-4o Vision to extract text from a PDF page rendered as an image.

    Advantages over Tesseract
    ─────────────────────────
    - Handles rotated, skewed, and low-quality scans
    - Understands document layout (headers, tables, signature blocks)
    - No dependency on language packs or Tesseract installation quality
    - Significantly better on Pakistani legal document typography

    Cost: ~USD 0.001 per page at GPT-4o Vision pricing (negligible for ~150 pages)
    """

    _SYSTEM = (
        "You are an OCR engine. Your only task is to output every character "
        "visible in the image exactly as it appears. Never summarise, interpret, "
        "explain, or refuse. Output raw text only."
    )

    _PROMPT = (
        "Output every word, number, and punctuation mark you can see in this "
        "image. Preserve line breaks and paragraph spacing. Do not add any "
        "commentary. If a region is blank, output nothing for that region."
    )

    # Phrases that indicate a refusal rather than extracted text
    _REFUSAL_PHRASES = (
        "i'm unable", "i cannot", "i can't", "i am unable",
        "unable to transcribe", "unable to provide", "i'm sorry",
        "i am sorry", "sorry, i", "cannot assist", "can't assist",
        "unable to assist", "not able to",
    )

    def __init__(self):
        from openai import OpenAI
        self._client = OpenAI(api_key=config.OPENAI_API_KEY)

    def ocr_page(self, image) -> str:
        """Extract text from a PIL image. Returns extracted text string.
        Raises ValueError if the model returns a refusal instead of text."""
        import base64
        import io

        buf = io.BytesIO()
        image.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        resp = self._client.chat.completions.create(
            model=config.VISION_OCR_MODEL,
            messages=[
                {"role": "system", "content": self._SYSTEM},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": self._PROMPT},
                        {"type": "image_url",
                         "image_url": {
                             "url":    f"data:image/png;base64,{b64}",
                             "detail": "high",
                         }},
                    ],
                },
            ],
            max_tokens=4096,
        )
        text = resp.choices[0].message.content.strip()

        # Detect refusals and raise so the caller falls back to Tesseract
        lower = text.lower()
        if any(phrase in lower for phrase in self._REFUSAL_PHRASES):
            raise ValueError(f"Vision OCR refused: {text[:120]!r}")

        return text


_vision_ocr: Optional[VisionOCR] = None

def _get_vision_ocr() -> Optional[VisionOCR]:
    """Lazy singleton — only instantiated when first needed."""
    global _vision_ocr
    if _vision_ocr is None and VISION_OCR_AVAILABLE:
        try:
            _vision_ocr = VisionOCR()
        except Exception:
            pass
    return _vision_ocr


# ══════════════════════════════════════════════════════════════════════════════
# Text extractors
# ══════════════════════════════════════════════════════════════════════════════

class PDFProcessor:
    """
    Extract text from PDF files.

    Strategy
    ────────
    1. Try pdfplumber (fast; works on text-layer PDFs).
    2. If a page yields < MIN_CHARS_PER_PAGE chars, fall back to Tesseract OCR
       via pdf2image (requires Tesseract + Poppler installed on the system).
    3. If OCR is not available, mark the page as scanned and continue.
    """

    # DPI for rasterising scanned pages before OCR
    OCR_DPI = 300

    def extract(self, path: Path) -> tuple[list[PageContent], list[str]]:
        pages: list[PageContent] = []
        notes: list[str] = []

        # ── Step 1: pdfplumber pass ───────────────────────────────────────────
        try:
            with pdfplumber.open(str(path)) as pdf:
                raw_pages = []
                for i, page in enumerate(pdf.pages, start=1):
                    text = (page.extract_text() or "").strip()
                    raw_pages.append((i, text))
        except Exception as exc:
            raise RuntimeError(f"pdfplumber failed on {path.name}: {exc}") from exc

        # ── Step 2: OCR fallback for sparse pages ─────────────────────────────
        scanned_page_indices = [
            i for i, (num, text) in enumerate(raw_pages)
            if len(text) < config.MIN_CHARS_PER_PAGE
        ]

        ocr_images: dict[int, object] = {}   # page_index → PIL image
        if scanned_page_indices:
            if not OCR_AVAILABLE:
                notes.append(
                    "OCR libraries (pytesseract/pdf2image) are available but "
                    "Tesseract is not installed. Install Tesseract OCR from "
                    "https://github.com/UB-Mannheim/tesseract/wiki and add it to PATH."
                    if OCR_AVAILABLE else
                    "Scanned pages detected. Install pytesseract + pdf2image + "
                    "Tesseract OCR to enable automatic text extraction."
                )
            else:
                try:
                    # Convert only the scanned page numbers (1-indexed) to images
                    page_numbers_to_ocr = [raw_pages[i][0] for i in scanned_page_indices]
                    images = convert_from_path(
                        str(path),
                        dpi=self.OCR_DPI,
                        first_page=min(page_numbers_to_ocr),
                        last_page=max(page_numbers_to_ocr),
                        poppler_path=config.POPPLER_PATH or None,
                    )
                    # Map images back to page indices
                    for img, page_num in zip(images, range(min(page_numbers_to_ocr),
                                                           max(page_numbers_to_ocr) + 1)):
                        idx = page_num - 1   # convert to 0-based
                        if idx in scanned_page_indices:
                            ocr_images[idx] = img
                    notes.append(
                        f"OCR applied to {len(ocr_images)} scanned page(s) "
                        f"({[raw_pages[i][0] for i in scanned_page_indices]})."
                    )
                except Exception as exc:
                    notes.append(
                        f"OCR conversion failed: {exc}. "
                        "Ensure Poppler is installed and in PATH (pdf2image requirement)."
                    )

        # ── Step 3: Assemble PageContent objects ──────────────────────────────
        vision = _get_vision_ocr()

        for idx, (page_num, plumber_text) in enumerate(raw_pages):
            if idx in ocr_images and len(plumber_text) < config.MIN_CHARS_PER_PAGE:

                # Run Tesseract first (always, when available) as the baseline
                tess_text = ""
                if OCR_AVAILABLE:
                    try:
                        tess_text = pytesseract.image_to_string(
                            ocr_images[idx], lang="eng", config="--psm 3"
                        ).strip()
                    except Exception as t_exc:
                        notes.append(f"Page {page_num}: Tesseract error – {t_exc}")

                # Run Vision OCR and keep whichever result is longer
                vision_text = ""
                if vision:
                    try:
                        vision_text = vision.ocr_page(ocr_images[idx])
                    except Exception as exc:
                        notes.append(f"Page {page_num}: Vision OCR failed ({exc}); "
                                     "using Tesseract result.")

                # Best-of-both: use whichever engine extracted more text
                if len(vision_text) >= len(tess_text):
                    ocr_text = vision_text
                    engine = f"Vision({len(vision_text)})"
                else:
                    ocr_text = tess_text
                    engine = f"Tesseract({len(tess_text)})"

                notes.append(f"Page {page_num}: {engine} chars — "
                             f"Vision={len(vision_text)} Tesseract={len(tess_text)} → kept {engine.split('(')[0]}.")

                final_text = ocr_text
                is_scanned = len(final_text) < config.MIN_CHARS_PER_PAGE
            else:
                final_text = plumber_text
                is_scanned = len(final_text) < config.MIN_CHARS_PER_PAGE
                if is_scanned and idx not in ocr_images:
                    notes.append(
                        f"Page {page_num}: {len(final_text)} chars extracted "
                        "(likely image-only; OCR required)"
                    )

            pages.append(PageContent(
                page_num=page_num,
                text=final_text,
                char_count=len(final_text),
                is_scanned=is_scanned,
            ))

        return pages, notes


class DocxProcessor:
    """Extract text from .docx files using python-docx."""

    def extract(self, path: Path) -> tuple[list[PageContent], list[str]]:
        notes: list[str] = []

        try:
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full = "\n".join(paragraphs)
        except Exception as exc:
            raise RuntimeError(f"python-docx failed on {path.name}: {exc}") from exc

        # Word docs have no physical page structure; treat as single page
        char_count = len(full)
        is_scanned = char_count < config.MIN_CHARS_PER_PAGE
        if is_scanned:
            notes.append("Very little text extracted from .docx – document may be image-only")

        pages = [PageContent(page_num=1, text=full, char_count=char_count, is_scanned=is_scanned)]
        return pages, notes


class DocProcessor:
    """.doc (legacy Word) – best-effort via python-docx; may fail on older binaries."""

    def extract(self, path: Path) -> tuple[list[PageContent], list[str]]:
        notes = ["Legacy .doc format: attempting python-docx extraction (may be incomplete)"]
        try:
            proc = DocxProcessor()
            return proc.extract(path)
        except Exception as exc:
            raise RuntimeError(
                f"Cannot extract .doc file '{path.name}'. "
                "Convert to .docx or install LibreOffice for conversion. "
                f"Detail: {exc}"
            ) from exc


# ══════════════════════════════════════════════════════════════════════════════
# Audit logger
# ══════════════════════════════════════════════════════════════════════════════

class AuditLogger:
    """
    Append-only JSONL audit trail.
    One record per ingestion attempt (success or failure).
    """

    def __init__(self, log_path: Path = config.AUDIT_LOG_PATH):
        self.log_path = log_path
        self.log_path.parent.mkdir(parents=True, exist_ok=True)

    def record(
        self,
        filename: str,
        doc_id: str,
        status: str,
        source_url: Optional[str],
        source_verified: bool,
        scanned_pages: int,
        total_pages: int,
        notes: list[str],
        error: Optional[str],
    ) -> None:
        entry = {
            "timestamp":       datetime.now(timezone.utc).isoformat(),
            "filename":        filename,
            "doc_id":          doc_id,
            "status":          status,
            "source_url":      source_url,
            "source_verified": source_verified,
            "total_pages":     total_pages,
            "scanned_pages":   scanned_pages,
            "notes":           notes,
            "error":           error,
        }
        with self.log_path.open("a", encoding="utf-8") as fh:
            fh.write(json.dumps(entry) + "\n")


# ══════════════════════════════════════════════════════════════════════════════
# Incremental tracker
# ══════════════════════════════════════════════════════════════════════════════

class IncrementalTracker:
    """
    Checks MongoDB to decide whether a file needs re-ingestion.
    A file is considered unchanged if a non-failed document with the same
    SHA-256 already exists in the documents collection.
    """

    @staticmethod
    def checksum(path: Path) -> str:
        h = hashlib.sha256()
        with path.open("rb") as fh:
            for chunk in iter(lambda: fh.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()

    def is_unchanged(self, path: Path) -> bool:
        """True if the file was already ingested with the same content."""
        import mongo_store
        return mongo_store.doc_exists_by_checksum(self.checksum(path))

    def mark(self, path: Path, checksum: str) -> None:
        pass  # MongoDB persistence is handled by upsert_doc


# ══════════════════════════════════════════════════════════════════════════════
# Main ingester
# ══════════════════════════════════════════════════════════════════════════════

class DocumentIngester:
    """
    Orchestrates the full Phase 1 pipeline for a single document or a folder.

    Usage
    ─────
    ingester = DocumentIngester()
    results  = ingester.ingest_folder(config.DOCS_DIR)
    """

    def __init__(self):
        self.validator  = DocumentValidator()
        self.pdf_proc   = PDFProcessor()
        self.docx_proc  = DocxProcessor()
        self.doc_proc   = DocProcessor()
        self.logger     = AuditLogger()
        self.tracker    = IncrementalTracker()

        config.PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
        config.FLAGGED_DIR.mkdir(parents=True, exist_ok=True)

    # ── public API ────────────────────────────────────────────────────────────

    def ingest_folder(
        self,
        folder: Path = config.DOCS_DIR,
        source_url_map: Optional[dict[str, str]] = None,
        force: bool = False,
    ) -> dict[str, str]:
        """
        Ingest all supported documents in *folder*.

        Parameters
        ──────────
        folder         : directory to scan
        source_url_map : optional {filename: "https://secp.gov.pk/..."} dict
        force          : re-ingest even if checksum matches (default False)

        Returns
        ───────
        {filename: status}  where status ∈ {success, partial, scanned, failed, skipped}
        """
        source_url_map = source_url_map or {}
        results: dict[str, str] = {}

        # Collect (path, category) pairs from category subfolders and root
        files: list[tuple[Path, str]] = []
        for entry in sorted(folder.iterdir()):
            if entry.is_dir():
                category = entry.name
                for p in sorted(entry.iterdir()):
                    if p.is_file() and p.suffix.lower() in config.SUPPORTED_EXTENSIONS:
                        files.append((p, category))
            elif entry.is_file() and entry.suffix.lower() in config.SUPPORTED_EXTENSIONS:
                files.append((entry, "Uncategorized"))

        print(f"\n{'-'*60}")
        print(f"  SECP Document Ingester  |  {len(files)} file(s) found in {folder.name}/")
        print(f"{'-'*60}")

        for path, category in files:
            url = source_url_map.get(path.name)
            status = self._ingest_one(path, url, force=force, category=category)
            results[path.name] = status
            icon = {"success": "OK", "partial": "~", "scanned": "!", "failed": "X", "skipped": "->"}.get(status, "?")
            print(f"  [{icon}] [{category[:30]:<30}] {path.name:<50}  {status}")

        counts = {s: sum(1 for v in results.values() if v == s)
                  for s in ("success", "partial", "scanned", "failed", "skipped")}
        print(f"\n  Summary: {counts}")
        print(f"{'-'*60}\n")
        return results

    def ingest_file(
        self,
        path: Path,
        source_url: Optional[str] = None,
        force: bool = False,
    ) -> str:
        """Ingest a single document. Returns status string."""
        return self._ingest_one(path, source_url, force=force)

    # ── internal pipeline ─────────────────────────────────────────────────────

    def _ingest_one(self, path: Path, source_url: Optional[str], force: bool, category: str = "Uncategorized") -> str:
        # ── 1. File existence check ───────────────────────────────────────────
        ok, msg = self.validator.validate_file_exists(path)
        if not ok:
            return self._fail(path, "", source_url, False, msg)

        # ── 2. Extension check ────────────────────────────────────────────────
        ok, msg = self.validator.validate_extension(path)
        if not ok:
            return self._fail(path, "", source_url, False, msg)

        # ── 3. Compute checksum; skip if unchanged ────────────────────────────
        checksum = self.tracker.checksum(path)
        if not force and self.tracker.is_unchanged(path):
            self.logger.record(
                filename=path.name, doc_id=checksum, status="skipped",
                source_url=source_url, source_verified=False,
                scanned_pages=0, total_pages=0, notes=["Checksum unchanged – skipped"], error=None,
            )
            return "skipped"

        # ── 4. Source URL validation ──────────────────────────────────────────
        source_verified, url_note = self.validator.validate_source_url(source_url)
        notes: list[str] = []
        if url_note:
            notes.append(url_note)

        # ── 5. Text extraction ────────────────────────────────────────────────
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                pages, extraction_notes = self.pdf_proc.extract(path)
            elif ext == ".docx":
                pages, extraction_notes = self.docx_proc.extract(path)
            else:  # .doc
                pages, extraction_notes = self.doc_proc.extract(path)
        except RuntimeError as exc:
            return self._fail(path, checksum, source_url, source_verified, str(exc))

        notes.extend(extraction_notes)

        # ── 6. Scanned-document detection ────────────────────────────────────
        total_pages   = len(pages)
        scanned_pages = sum(1 for p in pages if p.is_scanned)
        scanned_ratio = scanned_pages / total_pages if total_pages else 0

        if scanned_ratio >= config.SCANNED_DOC_RATIO:
            notes.append(
                f"{scanned_pages}/{total_pages} pages appear image-only "
                f"({scanned_ratio:.0%}). OCR required for full text extraction."
            )
            status = "scanned"
            self._copy_to_flagged(path)
        elif scanned_pages > 0:
            status = "partial"
        else:
            status = "success"

        # ── 7. Build full text ────────────────────────────────────────────────
        full_text = "\n\n".join(p.text for p in pages if p.text).strip()

        # ── 8. Parse filename metadata ────────────────────────────────────────
        order_meta = FilenameParser.parse(path.name)

        # ── 9. Assemble processed document ───────────────────────────────────
        doc = ProcessedDocument(
            doc_id=checksum,
            filename=path.name,
            file_path=str(path.resolve()),
            file_type=ext.lstrip("."),
            source_url=source_url,
            source_verified=source_verified,
            ingestion_timestamp=datetime.now(timezone.utc).isoformat(),
            status=status,
            order_metadata=order_meta,
            pages=pages,
            full_text=full_text,
            page_count=total_pages,
            scanned_page_count=scanned_pages,
            processing_notes=notes,
            error=None,
        )

        # ── 10. Persist to MongoDB ────────────────────────────────────────────
        import mongo_store
        serialised = self._serialise(doc)
        serialised["category"] = category
        mongo_store.upsert_doc(serialised)

        # ── 11. Update checksum store & audit log ─────────────────────────────
        self.tracker.mark(path, checksum)
        self.logger.record(
            filename=path.name, doc_id=checksum, status=status,
            source_url=source_url, source_verified=source_verified,
            scanned_pages=scanned_pages, total_pages=total_pages,
            notes=notes, error=None,
        )

        return status

    def _fail(
        self, path: Path, checksum: str,
        source_url: Optional[str], source_verified: bool, error: str,
    ) -> str:
        self._copy_to_flagged(path)
        self.logger.record(
            filename=path.name, doc_id=checksum or "unknown", status="failed",
            source_url=source_url, source_verified=source_verified,
            scanned_pages=0, total_pages=0, notes=[], error=error,
        )
        return "failed"

    @staticmethod
    def _copy_to_flagged(path: Path) -> None:
        dest = config.FLAGGED_DIR / path.name
        if not dest.exists():
            shutil.copy2(str(path), str(dest))

    @staticmethod
    def _serialise(doc: ProcessedDocument) -> dict:
        d = asdict(doc)
        # Trim heavy page.text from the top-level JSON – full text already in full_text
        # Keep page metadata (char_count, is_scanned) and first 500 chars as preview
        for page in d["pages"]:
            page["text_preview"] = page["text"][:500]
            del page["text"]
        return d


# ══════════════════════════════════════════════════════════════════════════════
# Audit log reader (utility)
# ══════════════════════════════════════════════════════════════════════════════

def read_audit_log(log_path: Path = config.AUDIT_LOG_PATH) -> list[dict]:
    """Return all audit log entries as a list of dicts."""
    if not log_path.exists():
        return []
    entries = []
    with log_path.open("r", encoding="utf-8") as fh:
        for line in fh:
            line = line.strip()
            if line:
                entries.append(json.loads(line))
    return entries


def print_audit_summary() -> None:
    entries = read_audit_log()
    if not entries:
        print("Audit log is empty.")
        return
    from collections import Counter
    counts = Counter(e["status"] for e in entries)
    unverified = sum(1 for e in entries if not e["source_verified"])
    flagged    = sum(1 for e in entries if e["status"] in ("failed", "scanned"))
    print(f"\nAudit Log Summary  ({len(entries)} total entries)")
    print(f"  Status breakdown : {dict(counts)}")
    print(f"  Unverified source: {unverified}")
    print(f"  Flagged for review: {flagged}\n")


# ══════════════════════════════════════════════════════════════════════════════
# Entry point
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="SECP Document Ingester – Phase 1")
    parser.add_argument(
        "--docs-dir", type=Path, default=config.DOCS_DIR,
        help="Folder containing SECP adjudication orders (default: SECP_docs/)",
    )
    parser.add_argument(
        "--force", action="store_true",
        help="Re-ingest all documents even if checksums match",
    )
    parser.add_argument(
        "--audit", action="store_true",
        help="Print audit log summary and exit",
    )
    parser.add_argument(
        "--url-map", type=Path, default=None,
        help="Optional JSON file mapping {filename: source_url} for domain verification",
    )
    parser.add_argument(
        "--reprocess-partial", action="store_true",
        help="Re-ingest only documents with status partial or scanned (uses Vision OCR)",
    )
    args = parser.parse_args()

    if args.audit:
        print_audit_summary()
    else:
        source_url_map: dict[str, str] = {}
        if args.url_map and args.url_map.exists():
            with args.url_map.open("r", encoding="utf-8") as fh:
                source_url_map = json.load(fh)

        ingester = DocumentIngester()

        if args.reprocess_partial:
            # Find docs with partial/scanned status from MongoDB
            import mongo_store
            partial_docs  = mongo_store.get_partial_docs()
            partial_files = [(Path(d["file_path"]), d.get("category", "Uncategorized"))
                             for d in partial_docs if d.get("file_path")]

            print(f"\n  Reprocessing {len(partial_files)} partial/scanned document(s) "
                  f"with Vision OCR...")
            vision = _get_vision_ocr()
            print(f"  Vision OCR : {'enabled (GPT-4o)' if vision else 'NOT available – check OPENAI_API_KEY'}")
            for path, category in partial_files:
                if path.exists():
                    status = ingester._ingest_one(path, None, force=True, category=category)
                    print(f"  [{status}] {path.name}")
                else:
                    print(f"  [!] File not found: {path}")
        else:
            ingester.ingest_folder(
                folder=args.docs_dir,
                source_url_map=source_url_map,
                force=args.force,
            )
        print_audit_summary()
